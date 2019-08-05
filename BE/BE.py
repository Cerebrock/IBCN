import docx
import os
import re

def get_files_that_endswith(end):
   paths = []
   for a, b, c in os.walk(os.getcwd()):
       for f in c:
           if f.endswith(end):
              path = os.path.abspath(f)
              paths.append(path)
   return paths


def quitar(format):
    for filename in get_files_that_endswith(format):
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        r = removeCor(fullText)
        rr = docx.Document()
        run = rr.add_paragraph().add_run(r)
        run.font.name = 'Arial'
        rr.save('%s SIN.docx' % filename)

def removeCor(fil):
    c = 0
    ret = []
    for l in fil:
        ref = True
        if len(l) > 5:
            for w in ['TEA', 'SD', 'EP', 'SPA']:
                if w in l:
                    ref = False
            if ref:
                ret.append(l + '\n\n')
        c += 1
    return ret

def contar(format):
    import pandas as pd
    df = pd.DataFrame(columns = ['N', 'TEA', 'SD', 'EP', 'SPA', 'Total'])
    for filename in get_files_that_endswith(format):
        n_p = re.findall('participante.+?(\d+)', filename)[0]
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        d = [w for w in re.findall('(TEA|SD|SPA|EP)[^:]', ''.join(fullText))]
        df = df.append(pd.concat([pd.Series(n_p, index= ['N']), pd.Series(d).value_counts(), pd.Series(len(d), index = ['Total'])]), ignore_index= True)
    df.to_excel('Data.xls')

inp = None
while inp not in ['0','1']:
    inp = input('Para quitar correcciones ingresar 0, para contarlas ingresar 1 \n')
    format = 'docx'
    #input('Tipo de archivo? \n')
    if inp == '0':
        quitar(format)
        exit()
    elif inp == '1':
        contar(format)
        exit()