import docx
import os

def get_files_that_endswith(end):
   paths = []
   for a, b, c in os.walk(os.getcwd()):
       for f in c:
           if f.endswith(end):
              path = os.path.abspath(f)
              paths.append(path)
   return paths


def do(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    r = removeCor(fullText)
    re = docx.Document()
    run = re.add_paragraph().add_run(r)
    run.font.name = 'Arial'
    re.save('%s SIN.docx' % filename)

def removeCor(fil):
    c = 0
    ret = []
    for l in fil:
        re = True
        if len(l) > 5:
            for w in ['TEA', 'SD', 'EP', 'SPA']:
                if w in l:
                    re = False
            if re:
                ret.append(l + '\n\n')
        c += 1
    return ret

for p in get_files_that_endswith('docx'):
    do(p)